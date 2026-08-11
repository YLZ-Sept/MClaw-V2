"""
知识库管理器（KnowledgeManager）

核心服务类，管理文档的完整生命周期：
- Collection 管理（创建/删除/列表）
- 文档摄入（解析 → 分块 → 嵌入 → 索引）
- 混合检索（BM25 + 向量 + RRF 融合）
- 优雅降级（ChromaDB 不可用时回退到纯 FTS5）

依赖：
- 复用 VectorStore 的 ChromaDB PersistentClient（独立 collection "knowledge"）
- 复用 model_hub 的 embedding 模型加载
- FTS5 + jieba 中文分词（零额外依赖）

Usage:
    manager = KnowledgeManager(data_dir=Path("data/knowledge"))
    coll = await manager.create_collection("技术文档", "项目相关技术资料")
    doc = await manager.ingest_file(coll.id, Path("doc.pdf"))
    results = await manager.search("如何使用", collection_id=coll.id, top_k=10)
"""

from __future__ import annotations

import asyncio
import hashlib
import logging
import os
import sqlite3
import threading
import time
from pathlib import Path
from typing import Any

from .knowledge_types import (
    CollectionStats,
    DocFileType,
    DocStatus,
    EXTENSION_MAP,
    KnowledgeCollection,
    KnowledgeChunk,
    KnowledgeDocument,
    SearchResult,
    _short_uuid,
)

logger = logging.getLogger(__name__)

# ── 延迟导入 ──────────────────────────────────────────────────────────────────
_chromadb = None
_sentence_transformers_available = False


def _lazy_import_chroma() -> bool:
    global _chromadb, _sentence_transformers_available
    if _chromadb is not None:
        return True
    try:
        from chromadb.config import Settings  # noqa: F401
        import chromadb as _c
        _chromadb = _c
        _sentence_transformers_available = True
        return True
    except ImportError:
        logger.info("[KnowledgeManager] ChromaDB 未安装，将使用纯 FTS5 模式")
        return False


# ── 文本提取器 ─────────────────────────────────────────────────────────────────

def _pdf_ocr_fallback(file_path: Path) -> str:
    """图片型 PDF 的 OCR 回退：pdf2image → pytesseract / paddleocr"""
    ocr_parts: list[str] = []

    # 尝试 pdf2image + pytesseract
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image

        images = convert_from_path(file_path, dpi=200)
        for i, img in enumerate(images):
            t = pytesseract.image_to_string(img, lang="chi_sim+eng")
            if t and t.strip():
                ocr_parts.append(t.strip())
        if ocr_parts:
            logger.info(f"[OCR] pytesseract 从 {len(images)} 页中提取了 {len(ocr_parts)} 页文字")
            return "\n\n".join(ocr_parts)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"[OCR] pytesseract 失败: {e}")

    # 回退: paddleocr（纯 Python，无需系统依赖）
    try:
        from paddleocr import PaddleOCR
        from pdf2image import convert_from_path

        ocr = PaddleOCR(lang="ch", use_angle_cls=True, show_log=False)
        images = convert_from_path(file_path, dpi=200)
        for img in images:
            import numpy as np
            result = ocr.ocr(np.array(img), cls=True)
            if result and result[0]:
                lines = [line[1][0] for line in result[0] if line[1][0]]
                if lines:
                    ocr_parts.append("\n".join(lines))
        if ocr_parts:
            logger.info(f"[OCR] PaddleOCR 从 {len(images)} 页中提取了 {len(ocr_parts)} 页文字")
            return "\n\n".join(ocr_parts)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"[OCR] PaddleOCR 失败: {e}")

    return ""


def _extract_doc_text(file_path: Path) -> str:
    """从旧版 .doc (Word 97-2003) 文件中提取文本"""
    # 方法 1: olefile 解析 OLE 二进制格式
    try:
        import olefile
        import struct

        ole = olefile.OleFileIO(str(file_path))
        if not ole.exists("WordDocument"):
            raise ValueError("不是有效的 .doc 文件：缺少 WordDocument 流")

        word_stream = ole.openstream("WordDocument").read()

        # 读取 FIB (File Information Block) 确定文本编码和位置
        # https://learn.microsoft.com/en-us/openspecs/office_file_formats/ms-doc/cc3f69a5
        flags = struct.unpack_from("<H", word_stream, 0x000A)[0]
        fComplex = bool(flags & 0x0200)

        if fComplex:
            # 复杂格式：文本存在 Table 流中
            # fcClx: file offset of Clx in table stream
            fcClx = struct.unpack_from("<I", word_stream, 0x01A2)[0]
            lcbClx = struct.unpack_from("<I", word_stream, 0x01A6)[0]

            table_stream_name = "1Table" if ole.exists("1Table") else "0Table"
            if ole.exists(table_stream_name):
                table_stream = ole.openstream(table_stream_name).read()
                clx = table_stream[fcClx:fcClx + lcbClx]

                # Prc 数据在最后一个 Clx 条目之后
                pos = 0
                while pos < len(clx):
                    entry_type = clx[pos]
                    if entry_type == 0x01:  # Prc
                        cb = struct.unpack_from("<H", clx, pos + 1)[0]
                        # 跳过 Prc 到达文本
                        text_offset = pos + 3 + cb
                        raw_text = clx[text_offset:]
                        ole.close()
                        # 尝试 Unicode 解码
                        try:
                            return raw_text.decode("utf-16-le", errors="ignore").replace("\x00", "").strip()
                        except Exception:
                            return raw_text.decode("cp1252", errors="ignore").strip()
                    elif entry_type == 0x02:  # Pcdt
                        _lcb = struct.unpack_from("<I", clx, pos + 1)[0]
                        pos += 5 + _lcb
                    else:
                        break
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"[DOC] olefile 解析失败: {e}")

    # 方法 2: python-docx 对 .doc 不支持，直接尝试读裸字节
    try:
        raw = file_path.read_bytes()
        # 尝试 UTF-16LE
        try:
            text = raw.decode("utf-16-le", errors="ignore")
            # 过滤掉过多的不可打印字符
            printable = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            if len(printable) > 100:
                return printable.strip()
        except Exception:
            pass
        # 尝试 ASCII / Latin-1
        text = raw.decode("latin-1", errors="ignore")
        import re
        words = re.findall(r"[\x20-\x7E一-鿿　-〿＀-￯]{3,}", text)
        if words:
            return " ".join(words)
    except Exception as e:
        logger.warning(f"[DOC] 原始读取失败: {e}")

    raise ImportError(
        "无法解析 .doc 文件。请安装 olefile 并将文件转为 .docx 格式后重试: "
        "pip install olefile"
    )


def _extract_text(file_path: Path, file_type: DocFileType) -> str:
    """从文件中提取纯文本"""
    if file_type == DocFileType.TEXT or file_type == DocFileType.MARKDOWN:
        return file_path.read_text(encoding="utf-8", errors="replace")

    if file_type == DocFileType.PDF:
        text = ""
        # 第 1 步: pdfplumber（文字型 PDF）
        try:
            import pdfplumber
            text_parts: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            text = "\n\n".join(text_parts)
        except ImportError:
            # fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(file_path))
                text = "\n\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
            except ImportError:
                raise ImportError(
                    "PDF 解析需要安装 pdfplumber 或 PyPDF2: "
                    "pip install pdfplumber"
                )

        # 第 2 步: 文字提取为空 → 可能是图片型 PDF，尝试 OCR
        if not text.strip():
            ocr_text = _pdf_ocr_fallback(file_path)
            if ocr_text:
                return ocr_text

        return text

    if file_type == DocFileType.DOCX:
        try:
            from docx import Document
            doc = Document(str(file_path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ImportError(
                "DOCX 解析需要安装 python-docx: pip install python-docx"
            )

    if file_type == DocFileType.DOC:
        return _extract_doc_text(file_path)

    if file_type == DocFileType.HTML:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_path.read_text(encoding="utf-8", errors="replace"), "html.parser")
            # Remove script/style tags
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # bare fallback: strip tags naively
            import re
            html = file_path.read_text(encoding="utf-8", errors="replace")
            text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<[^>]+>", "", text)
            return text.strip()

    if file_type in (DocFileType.XLSX, DocFileType.XLS):
        return _extract_excel(file_path, file_type)

    if file_type == DocFileType.CSV:
        return file_path.read_text(encoding="utf-8", errors="replace")

    raise ValueError(f"不支持的文件类型: {file_type}")


def _extract_excel(file_path: Path, file_type: DocFileType) -> str:
    """从 Excel 文件提取文本（每个单元格一行，sheet 间加分隔）。"""
    if file_type == DocFileType.XLSX:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            parts: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        parts.append(row_text)
            wb.close()
            return "\n".join(parts)
        except ImportError:
            raise ImportError(
                "Excel(.xlsx) 解析需要安装 openpyxl: pip install openpyxl"
            )
    else:
        # .xls — older format
        try:
            import xlrd
            wb = xlrd.open_workbook(str(file_path))
            parts: list[str] = []
            for sheet_name in wb.sheet_names():
                ws = wb.sheet_by_name(sheet_name)
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row_idx in range(ws.nrows):
                    row_text = "\t".join(
                        str(ws.cell_value(row_idx, col_idx))
                        for col_idx in range(ws.ncols)
                    )
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        except ImportError:
            raise ImportError(
                "Excel(.xls) 解析需要安装 xlrd: pip install xlrd"
            )


def _extract_text_from_url(url: str) -> str:
    """从 URL 下载并提取文本（尝试 requests + readability）"""
    import re

    try:
        import requests
        resp = requests.get(url, timeout=30, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        resp.raise_for_status()
        html = resp.text
        # If encoding is not set, try to detect
        if resp.encoding and resp.encoding.lower() != "utf-8":
            html = resp.content.decode(resp.encoding, errors="replace")
    except ImportError:
        raise ImportError("URL 导入需要安装 requests: pip install requests")

    # Try readability-lxml for main content extraction
    try:
        from readability import Document
        doc = Document(html)
        title = doc.title()
        content_html = doc.summary()
        # Strip tags from summary
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content_html, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        return f"# {title}\n\n{text}" if title else text
    except ImportError:
        pass

    # Fallback: BeautifulSoup
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "", text)
        return text.strip()


# ── 分块器 ────────────────────────────────────────────────────────────────────

def _chunk_text(
    text: str,
    chunk_size: int = 512,
    chunk_overlap: int = 128,
) -> list[str]:
    """滑动窗口分块，按段落优先分割"""
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= chunk_size:
            current = f"{current}\n\n{para}" if current else para
        else:
            if current:
                chunks.append(current)
            # If a single paragraph exceeds chunk_size, split by sentence
            if len(para) > chunk_size:
                sentences = _split_sentences(para)
                buf = ""
                for s in sentences:
                    if len(buf) + len(s) <= chunk_size:
                        buf = f"{buf}{s}" if buf else s
                    else:
                        if buf:
                            chunks.append(buf)
                        buf = s
                if buf:
                    current = buf
                else:
                    current = ""
            else:
                current = para

    if current:
        chunks.append(current)

    # Apply overlap: each chunk (except first) prepends overlap chars from previous
    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            if len(prev) > chunk_overlap:
                overlapped.append(prev[-chunk_overlap:] + "\n\n" + chunks[i])
            else:
                overlapped.append(chunks[i])
        chunks = overlapped

    return chunks


def _split_sentences(text: str) -> list[str]:
    """中英文分句"""
    import re
    # Split on Chinese/English sentence endings
    parts = re.split(r"(?<=[。！？.!?\n])\s*", text)
    return [p.strip() for p in parts if p.strip()]


def _estimate_tokens(text: str) -> int:
    """粗略估计 token 数（中文按字符，英文按空格分词）"""
    import re
    chinese_chars = len(re.findall(r"[一-鿿]", text))
    english_words = len(re.findall(r"[a-zA-Z]+", text))
    others = len(text) - chinese_chars - english_words
    return chinese_chars + english_words + others // 4


# ── KnowledgeManager ───────────────────────────────────────────────────────────

class KnowledgeManager:
    """
    知识库管理器

    管理文档摄入、存储和混合检索。内部使用 SQLite 存储元数据，
    ChromaDB 存储向量（可降级到纯 FTS5）。
    """

    def __init__(
        self,
        data_dir: Path,
        model_name: str | None = None,
        device: str = "cpu",
        download_source: str = "auto",
    ):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.model_name = model_name or "shibing624/text2vec-base-chinese"
        self.device = device
        self.download_source = download_source

        # ChromaDB state
        self._chroma_client = None
        self._chroma_collection = None
        self._embedding_model = None
        self._vector_enabled = False
        self._init_lock = threading.RLock()
        self._init_started = False

        # SQLite database
        self._db_path = self.data_dir / "knowledge.db"
        self._init_sqlite()

        # Start background ChromaDB init
        self._start_chroma_init()

    # ── SQLite ──────────────────────────────────────────────────────────────

    @staticmethod
    def _tokenize_for_fts(text: str) -> str:
        """Pre-tokenize text for FTS5 indexing.

        Splits CJK characters into space-separated bigrams so the unicode61
        tokenizer can index and match them. Non-CJK text passes through.
        """
        try:
            import jieba
            jieba.setLogLevel(logging.WARNING)
            return " ".join(jieba.cut_for_search(text))
        except ImportError:
            pass
        # Fallback: insert spaces between CJK characters
        import re
        result = re.sub(r"([一-鿿㐀-䶿])", r" \1 ", text)
        return re.sub(r"\s+", " ", result).strip()

    def _init_sqlite(self) -> None:
        """创建 SQLite 表结构"""
        conn = sqlite3.connect(str(self._db_path))
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA foreign_keys=ON")
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS kb_collections (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT DEFAULT '',
                workspace_id TEXT DEFAULT 'default',
                owner_id TEXT DEFAULT '',
                is_public INTEGER DEFAULT 0,
                created_at REAL NOT NULL,
                updated_at REAL NOT NULL
            );

            CREATE TABLE IF NOT EXISTS kb_documents (
                id TEXT PRIMARY KEY,
                collection_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                file_path TEXT,
                file_type TEXT NOT NULL,
                file_size INTEGER DEFAULT 0,
                checksum TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                chunk_count INTEGER DEFAULT 0,
                error_message TEXT,
                uploaded_by TEXT DEFAULT '',
                created_at REAL NOT NULL,
                updated_at REAL NOT NULL,
                metadata_json TEXT DEFAULT '{}',
                FOREIGN KEY (collection_id) REFERENCES kb_collections(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS kb_chunks (
                id TEXT PRIMARY KEY,
                doc_id TEXT NOT NULL,
                collection_id TEXT NOT NULL,
                content TEXT NOT NULL,
                chunk_index INTEGER DEFAULT 0,
                token_count INTEGER DEFAULT 0,
                embedding_id TEXT,
                metadata_json TEXT DEFAULT '{}',
                FOREIGN KEY (doc_id) REFERENCES kb_documents(id) ON DELETE CASCADE,
                FOREIGN KEY (collection_id) REFERENCES kb_collections(id) ON DELETE CASCADE
            );

            CREATE VIRTUAL TABLE IF NOT EXISTS kb_chunks_fts USING fts5(
                content,
                tokenize='unicode61'
            );

            CREATE INDEX IF NOT EXISTS idx_kb_docs_collection
                ON kb_documents(collection_id);
            CREATE INDEX IF NOT EXISTS idx_kb_chunks_doc
                ON kb_chunks(doc_id);
            CREATE INDEX IF NOT EXISTS idx_kb_chunks_collection
                ON kb_chunks(collection_id);
        """)
        # Migration: add uploaded_by to existing tables
        try:
            conn.execute("ALTER TABLE kb_documents ADD COLUMN uploaded_by TEXT DEFAULT ''")
        except sqlite3.OperationalError:
            pass  # column already exists
        # Migration: add workspace/owner/public to collections (P1 multi-tenancy)
        for col, default in [("workspace_id", "'default'"), ("owner_id", "''"), ("is_public", "0")]:
            try:
                conn.execute(f"ALTER TABLE kb_collections ADD COLUMN {col} TEXT DEFAULT {default}")
            except sqlite3.OperationalError:
                pass
        conn.commit()
        conn.close()
        logger.info(f"[KnowledgeManager] SQLite 初始化完成: {self._db_path}")

    def _get_conn(self) -> sqlite3.Connection:
        conn = sqlite3.connect(str(self._db_path))
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys=ON")
        return conn

    # ── ChromaDB 初始化 ─────────────────────────────────────────────────────

    def _start_chroma_init(self) -> None:
        """后台线程初始化 ChromaDB（不阻塞）"""
        with self._init_lock:
            if self._init_started:
                return
            self._init_started = True

        t = threading.Thread(
            target=self._do_chroma_init,
            name="KnowledgeManager-chroma-init",
            daemon=True,
        )
        t.start()

    def _do_chroma_init(self) -> None:
        """实际 ChromaDB 初始化逻辑"""
        if not _lazy_import_chroma():
            logger.info("[KnowledgeManager] ChromaDB 不可用，使用纯 FTS5 模式")
            return

        try:
            # 加载 embedding 模型
            from .model_hub import load_embedding_model, _apply_source_env, _resolve_source, detect_best_source

            resolved = _resolve_source(self.download_source)
            if resolved.value == "auto":
                resolved = detect_best_source()
            _apply_source_env(resolved)

            logger.info(f"[KnowledgeManager] 加载 embedding 模型: {self.model_name}")
            model = load_embedding_model(
                model_name=self.model_name,
                source=self.download_source,
                device=self.device,
            )

            # 初始化 ChromaDB client
            chromadb_dir = self.data_dir / "chromadb"
            chromadb_dir.mkdir(parents=True, exist_ok=True)

            from chromadb.config import Settings

            client = _chromadb.PersistentClient(
                path=str(chromadb_dir),
                settings=Settings(anonymized_telemetry=False),
            )

            # 创建独立的 knowledge collection
            collection = client.get_or_create_collection(
                name="knowledge",
                metadata={"hnsw:space": "cosine"},
            )

            with self._init_lock:
                self._embedding_model = model
                self._chroma_client = client
                self._chroma_collection = collection
                self._vector_enabled = True

            logger.info(
                f"[KnowledgeManager] ChromaDB ✓ 已就绪 "
                f"(collection='knowledge', {collection.count()} vectors)"
            )

        except Exception as e:
            logger.warning(
                f"[KnowledgeManager] ChromaDB 初始化失败，使用纯 FTS5 模式: {e}",
                exc_info=True,
            )

    # ── Collection CRUD ─────────────────────────────────────────────────────

    def create_collection(self, name: str, description: str = "",
                          workspace_id: str = "default", owner_id: str = "",
                          is_public: bool = False) -> KnowledgeCollection:
        now = time.time()
        coll = KnowledgeCollection(
            name=name,
            description=description,
            workspace_id=workspace_id,
            owner_id=owner_id,
            is_public=is_public,
            created_at=now,
            updated_at=now,
        )
        conn = self._get_conn()
        conn.execute(
            "INSERT INTO kb_collections (id, name, description, workspace_id, "
            "owner_id, is_public, created_at, updated_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (coll.id, coll.name, coll.description, coll.workspace_id,
             coll.owner_id, int(coll.is_public), coll.created_at, coll.updated_at),
        )
        conn.commit()
        conn.close()
        logger.info(f"[KnowledgeManager] 创建 collection: {coll.name} ({coll.id}) ws={workspace_id}")
        return coll

    def list_collections(self, workspace_id: str | None = None) -> list[KnowledgeCollection]:
        conn = self._get_conn()
        if workspace_id:
            rows = conn.execute(
                "SELECT c.*, "
                "  (SELECT COUNT(*) FROM kb_documents WHERE collection_id = c.id) AS doc_count, "
                "  (SELECT COUNT(*) FROM kb_chunks WHERE collection_id = c.id) AS chunk_count "
                "FROM kb_collections c WHERE c.workspace_id = ? ORDER BY c.updated_at DESC",
                (workspace_id,),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT c.*, "
                "  (SELECT COUNT(*) FROM kb_documents WHERE collection_id = c.id) AS doc_count, "
                "  (SELECT COUNT(*) FROM kb_chunks WHERE collection_id = c.id) AS chunk_count "
                "FROM kb_collections c ORDER BY c.updated_at DESC"
            ).fetchall()
        conn.close()
        # sqlite3.Row 不支持 .get()，用 dict(r) 转换
        return [
            KnowledgeCollection(
                id=r["id"], name=r["name"], description=r["description"],
                workspace_id=dict(r).get("workspace_id", "default"),
                owner_id=dict(r).get("owner_id", ""),
                is_public=bool(dict(r).get("is_public", 0)),
                created_at=r["created_at"], updated_at=r["updated_at"],
                doc_count=r["doc_count"], chunk_count=r["chunk_count"],
            )
            for r in rows
        ]

    def get_collection(self, collection_id: str) -> KnowledgeCollection | None:
        conn = self._get_conn()
        r = conn.execute(
            "SELECT c.*, "
            "  (SELECT COUNT(*) FROM kb_documents WHERE collection_id = c.id) AS doc_count, "
            "  (SELECT COUNT(*) FROM kb_chunks WHERE collection_id = c.id) AS chunk_count "
            "FROM kb_collections c WHERE c.id = ?",
            (collection_id,),
        ).fetchone()
        conn.close()
        if r is None:
            return None
        d = dict(r)
        return KnowledgeCollection(
            id=d["id"], name=d["name"], description=d["description"],
            workspace_id=d.get("workspace_id", "default"),
            owner_id=d.get("owner_id", ""),
            is_public=bool(d.get("is_public", 0)),
            created_at=d["created_at"], updated_at=d["updated_at"],
            doc_count=d["doc_count"], chunk_count=d["chunk_count"],
        )

    def update_collection(self, collection_id: str, name: str | None = None, description: str | None = None) -> bool:
        updates: list[str] = []
        params: list[Any] = []
        if name is not None:
            updates.append("name = ?")
            params.append(name)
        if description is not None:
            updates.append("description = ?")
            params.append(description)
        if not updates:
            return False
        updates.append("updated_at = ?")
        params.append(time.time())
        params.append(collection_id)

        conn = self._get_conn()
        conn.execute(f"UPDATE kb_collections SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()
        conn.close()
        return True

    def delete_collection(self, collection_id: str) -> int:
        """删除 collection 及其所有文档、chunks、向量，返回删除的 chunk 数"""
        conn = self._get_conn()
        # 统计 chunk 数（用于向量删除）
        chunk_count = conn.execute(
            "SELECT COUNT(*) FROM kb_chunks WHERE collection_id = ?", (collection_id,)
        ).fetchone()[0]

        # 收集 embedding_ids 用于删除向量
        emb_rows = conn.execute(
            "SELECT embedding_id FROM kb_chunks WHERE collection_id = ? AND embedding_id IS NOT NULL",
            (collection_id,),
        ).fetchall()
        emb_ids = [r["embedding_id"] for r in emb_rows if r["embedding_id"]]

        # Cascade delete: documents + chunks (FTS5 trigger handles kb_chunks_fts)
        conn.execute("DELETE FROM kb_chunks WHERE collection_id = ?", (collection_id,))
        conn.execute("DELETE FROM kb_documents WHERE collection_id = ?", (collection_id,))
        conn.execute("DELETE FROM kb_chunks_fts WHERE rowid NOT IN (SELECT rowid FROM kb_chunks)")
        conn.execute("DELETE FROM kb_collections WHERE id = ?", (collection_id,))
        conn.commit()
        conn.close()

        # 删除 ChromaDB 向量
        if emb_ids and self._chroma_collection:
            try:
                self._chroma_collection.delete(ids=emb_ids)
            except Exception as e:
                logger.warning(f"[KnowledgeManager] 删除向量失败: {e}")

        logger.info(f"[KnowledgeManager] 已删除 collection={collection_id} ({chunk_count} chunks)")
        return chunk_count

    def get_stats(self, collection_id: str) -> CollectionStats:
        conn = self._get_conn()
        # Use subqueries to avoid cross-join between documents and chunks
        r = conn.execute(
            "SELECT "
            "  (SELECT COUNT(*) FROM kb_documents WHERE collection_id = ?) AS doc_count, "
            "  (SELECT COUNT(*) FROM kb_chunks WHERE collection_id = ?) AS chunk_count, "
            "  (SELECT COALESCE(SUM(LENGTH(content)), 0) FROM kb_chunks WHERE collection_id = ?) AS total_chars, "
            "  (SELECT COALESCE(SUM(file_size), 0) FROM kb_documents WHERE collection_id = ?) AS storage_bytes, "
            "  (SELECT COUNT(*) FROM kb_chunks WHERE collection_id = ? AND embedding_id IS NOT NULL) AS indexed_count, "
            "  (SELECT COUNT(*) FROM kb_documents WHERE collection_id = ? AND status = 'error') AS error_count",
            (collection_id,) * 6,
        ).fetchone()
        conn.close()
        return CollectionStats(
            collection_id=collection_id,
            doc_count=r["doc_count"] or 0,
            chunk_count=r["chunk_count"] or 0,
            total_chars=r["total_chars"] or 0,
            storage_bytes=r["storage_bytes"] or 0,
            indexed_count=r["indexed_count"] or 0,
            error_count=r["error_count"] or 0,
        )

    # ── Document CRUD ───────────────────────────────────────────────────────

    def list_documents(self, collection_id: str) -> list[KnowledgeDocument]:
        conn = self._get_conn()
        rows = conn.execute(
            "SELECT * FROM kb_documents WHERE collection_id = ? ORDER BY created_at DESC",
            (collection_id,),
        ).fetchall()
        conn.close()
        return [_doc_from_row(r) for r in rows]

    def get_document(self, doc_id: str) -> KnowledgeDocument | None:
        conn = self._get_conn()
        r = conn.execute("SELECT * FROM kb_documents WHERE id = ?", (doc_id,)).fetchone()
        conn.close()
        return _doc_from_row(r) if r else None

    def get_document_chunks(self, doc_id: str) -> list[KnowledgeChunk]:
        conn = self._get_conn()
        rows = conn.execute(
            "SELECT * FROM kb_chunks WHERE doc_id = ? ORDER BY chunk_index",
            (doc_id,),
        ).fetchall()
        conn.close()
        return [_chunk_from_row(r) for r in rows]

    def delete_document(self, doc_id: str) -> int:
        """删除文档及其 chunks 和向量"""
        conn = self._get_conn()
        doc = conn.execute("SELECT * FROM kb_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            conn.close()
            return 0

        # 收集 embedding_ids
        emb_rows = conn.execute(
            "SELECT embedding_id FROM kb_chunks WHERE doc_id = ? AND embedding_id IS NOT NULL",
            (doc_id,),
        ).fetchall()
        emb_ids = [r["embedding_id"] for r in emb_rows]

        # 删除
        chunk_count = conn.execute("DELETE FROM kb_chunks WHERE doc_id = ?", (doc_id,)).rowcount
        conn.execute("DELETE FROM kb_chunks_fts WHERE rowid NOT IN (SELECT rowid FROM kb_chunks)")
        conn.execute("DELETE FROM kb_documents WHERE id = ?", (doc_id,))
        conn.commit()
        conn.close()

        # 删除向量
        if emb_ids and self._chroma_collection:
            try:
                self._chroma_collection.delete(ids=emb_ids)
            except Exception as e:
                logger.warning(f"[KnowledgeManager] 删除向量失败: {e}")

        logger.info(f"[KnowledgeManager] 已删除文档 {doc_id} ({chunk_count} chunks)")
        return chunk_count

    # ── 文档摄入 ────────────────────────────────────────────────────────────

    def ingest_file(self, collection_id: str, file_path: Path,
                    uploaded_by: str = "") -> KnowledgeDocument:
        """摄入单个文件：解析 → 分块 → 索引"""
        suffix = file_path.suffix.lower()
        if suffix not in EXTENSION_MAP:
            raise ValueError(f"不支持的文件格式: {suffix}")

        file_type = EXTENSION_MAP[suffix]
        file_size = file_path.stat().st_size
        checksum = _sha256_file(file_path)

        # 检查是否已存在（相同路径+内容）
        conn = self._get_conn()
        existing = conn.execute(
            "SELECT id, checksum FROM kb_documents WHERE collection_id = ? AND file_path = ?",
            (collection_id, str(file_path)),
        ).fetchone()
        if existing and existing["checksum"] == checksum:
            conn.close()
            logger.info(f"[KnowledgeManager] 文档未变更，跳过: {file_path.name}")
            return _doc_from_row(existing)

        # 创建文档记录
        now = time.time()
        doc = KnowledgeDocument(
            collection_id=collection_id,
            filename=file_path.name,
            file_path=str(file_path),
            file_type=file_type.value,
            file_size=file_size,
            checksum=checksum,
            status=DocStatus.PENDING.value,
            uploaded_by=uploaded_by,
            created_at=now,
            updated_at=now,
        )

        if existing:
            # 更新已有记录
            doc.id = existing["id"]
            conn.execute(
                "UPDATE kb_documents SET checksum=?, status='pending', file_size=?, "
                "updated_at=? WHERE id=?",
                (checksum, file_size, now, doc.id),
            )
            # 删除旧 chunks
            conn.execute("DELETE FROM kb_chunks WHERE doc_id = ?", (doc.id,))
        else:
            conn.execute(
                "INSERT INTO kb_documents (id, collection_id, filename, file_path, "
                "file_type, file_size, checksum, status, chunk_count, uploaded_by, created_at, updated_at) "
                "VALUES (?,?,?,?,?,?,?,?,0,?,?,?)",
                (doc.id, doc.collection_id, doc.filename, doc.file_path,
                 doc.file_type, doc.file_size, doc.checksum, doc.status,
                 doc.uploaded_by, doc.created_at, doc.updated_at),
            )

        conn.commit()
        conn.close()

        # 执行索引
        try:
            self._index_document(doc, file_path, file_type)
        except Exception as e:
            self._mark_doc_error(doc.id, str(e))
            raise

        return self.get_document(doc.id) or doc

    def ingest_url(self, collection_id: str, url: str,
                   uploaded_by: str = "") -> KnowledgeDocument:
        """从 URL 摄入内容"""
        import hashlib as _hashlib

        now = time.time()
        url_hash = _hashlib.sha256(url.encode()).hexdigest()[:16]
        text = _extract_text_from_url(url)
        checksum = _hashlib.sha256(text.encode()).hexdigest()

        doc = KnowledgeDocument(
            collection_id=collection_id,
            filename=f"URL: {url[:80]}",
            file_type=DocFileType.URL.value,
            file_size=len(text.encode()),
            checksum=checksum,
            status=DocStatus.INDEXING.value,
            uploaded_by=uploaded_by,
            created_at=now,
            updated_at=now,
            metadata={"source_url": url, "url_hash": url_hash},
        )

        conn = self._get_conn()
        conn.execute(
            "INSERT INTO kb_documents (id, collection_id, filename, file_path, "
            "file_type, file_size, checksum, status, chunk_count, created_at, updated_at, metadata_json) "
            "VALUES (?,?,?,?,?,?,?,?,0,?,?,?)",
            (doc.id, doc.collection_id, doc.filename, None,
             doc.file_type, doc.file_size, doc.checksum, doc.status,
             doc.created_at, doc.updated_at, str(doc.metadata).replace("'", '"')),
        )
        conn.commit()
        conn.close()

        # Index chunks from extracted text
        chunks = _chunk_text(text)
        self._persist_chunks(doc, chunks)
        self._embed_chunks(doc, chunks)

        self._mark_doc_ready(doc.id, len(chunks))
        return self.get_document(doc.id) or doc

    def _index_document(self, doc: KnowledgeDocument, file_path: Path, file_type: DocFileType) -> None:
        """执行实际的文档索引"""
        self._mark_doc_status(doc.id, DocStatus.INDEXING)

        text = _extract_text(file_path, file_type)
        chunks = _chunk_text(text)

        if not chunks:
            suffix = file_path.suffix.lower()
            hint = ""
            if suffix == ".pdf":
                hint = (
                    "（PDF 可能为图片扫描件，文字提取为空。"
                    "请安装 Tesseract OCR 和 pytesseract："
                    "https://github.com/UB-Mannheim/tesseract/wiki）"
                )
            raise ValueError(f"文档内容为空: {file_path.name}{hint}")

        self._persist_chunks(doc, chunks)
        self._embed_chunks(doc, chunks)
        self._mark_doc_ready(doc.id, len(chunks))

        logger.info(
            f"[KnowledgeManager] 索引完成: {doc.filename} → {len(chunks)} chunks"
        )

    def _persist_chunks(self, doc: KnowledgeDocument, chunks: list[str]) -> None:
        """将 chunks 写入 SQLite + FTS5"""
        conn = self._get_conn()
        for i, text in enumerate(chunks):
            chunk_id = _short_uuid()
            token_count = _estimate_tokens(text)
            conn.execute(
                "INSERT INTO kb_chunks (id, doc_id, collection_id, content, chunk_index, token_count) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (chunk_id, doc.id, doc.collection_id, text, i, token_count),
            )
            # 同步写入 FTS5（预分词）
            tokenized = self._tokenize_for_fts(text)
            conn.execute(
                "INSERT INTO kb_chunks_fts (rowid, content) VALUES ("
                "  (SELECT rowid FROM kb_chunks WHERE id = ?), ?"
                ")",
                (chunk_id, tokenized),
            )
        conn.commit()
        conn.close()

    def _embed_chunks(self, doc: KnowledgeDocument, chunks: list[str]) -> None:
        """向量化 chunks 并存入 ChromaDB"""
        if not self._vector_enabled or not self._embedding_model:
            logger.info("[KnowledgeManager] 向量未启用，跳过 embedding")
            return

        # 读取刚写入的 chunk ids
        conn = self._get_conn()
        chunk_rows = conn.execute(
            "SELECT id, content FROM kb_chunks WHERE doc_id = ? ORDER BY chunk_index",
            (doc.id,),
        ).fetchall()

        texts = [r["content"] for r in chunk_rows]
        chunk_ids = [r["id"] for r in chunk_rows]
        embedding_ids = [f"kb_{cid}" for cid in chunk_ids]

        try:
            # 生成 embeddings
            embeddings = self._embedding_model.encode(
                texts,
                show_progress_bar=False,
                batch_size=32,
            ).tolist()

            # 存入 ChromaDB
            self._chroma_collection.add(
                ids=embedding_ids,
                embeddings=embeddings,
                metadatas=[
                    {"chunk_id": cid, "doc_id": doc.id, "collection_id": doc.collection_id}
                    for cid in chunk_ids
                ],
            )

            # 更新 embedding_id
            for cid, eid in zip(chunk_ids, embedding_ids):
                conn.execute(
                    "UPDATE kb_chunks SET embedding_id = ? WHERE id = ?",
                    (eid, cid),
                )
            conn.commit()

            logger.info(
                f"[KnowledgeManager] 向量化完成: {doc.filename} → {len(chunks)} vectors"
            )

        except Exception as e:
            logger.warning(f"[KnowledgeManager] 向量化失败: {e}")
        finally:
            conn.close()

    def reindex_document(self, doc_id: str) -> KnowledgeDocument:
        """重新索引文档"""
        doc = self.get_document(doc_id)
        if not doc:
            raise ValueError(f"文档不存在: {doc_id}")

        if doc.file_type == DocFileType.URL.value:
            url = doc.metadata.get("source_url", "")
            if not url:
                raise ValueError("URL 文档缺少 source_url")
            text = _extract_text_from_url(url)
        elif doc.file_path:
            file_path = Path(doc.file_path)
            if not file_path.exists():
                raise FileNotFoundError(
                    f"文档文件已丢失: {file_path.name}，请删除此记录后重新上传"
                )
            file_type = DocFileType(doc.file_type) if doc.file_type in [e.value for e in DocFileType] else DocFileType.TEXT
            text = _extract_text(file_path, file_type)
        else:
            raise ValueError(f"无法定位文档内容: {doc_id}")

        # 删除旧 chunks 和向量
        old_emb = self._get_chunk_embedding_ids(doc_id)
        conn = self._get_conn()
        conn.execute("DELETE FROM kb_chunks WHERE doc_id = ?", (doc_id,))
        conn.execute("DELETE FROM kb_chunks_fts WHERE rowid NOT IN (SELECT rowid FROM kb_chunks)")
        conn.commit()
        conn.close()
        if old_emb and self._chroma_collection:
            try:
                self._chroma_collection.delete(ids=old_emb)
            except Exception:
                pass

        # 重新分块索引
        chunks = _chunk_text(text)
        self._mark_doc_status(doc_id, DocStatus.INDEXING)
        self._persist_chunks(doc, chunks)
        self._embed_chunks(doc, chunks)
        self._mark_doc_ready(doc_id, len(chunks))

        return self.get_document(doc_id) or doc

    # ── 混合检索 ────────────────────────────────────────────────────────────

    def search(
        self,
        query: str,
        collection_id: str | None = None,
        top_k: int = 10,
    ) -> list[SearchResult]:
        """
        混合检索：BM25(FTS5+Jieba) + 向量(ChromaDB) → RRF 融合

        Args:
            query: 搜索查询
            collection_id: 可选，限定 collection
            top_k: 返回结果数
        """
        bm25_results = self._search_bm25(query, collection_id, top_k * 2)
        vector_results = self._search_vector(query, collection_id, top_k * 2)
        merged = self._rrf_fusion(bm25_results, vector_results, k=60)
        merged.sort(key=lambda x: x.score, reverse=True)
        return merged[:top_k]

    @staticmethod
    def _sanitize_fts_query(query: str) -> str:
        """Make user input safe for FTS5 MATCH."""
        special = set('"*(){}[]^~:')
        cleaned = "".join(c if c not in special else " " for c in query)
        tokens = cleaned.split()
        if not tokens:
            return '""'
        return " OR ".join(tokens)

    def _search_bm25(
        self, query: str, collection_id: str | None, limit: int
    ) -> dict[str, float]:
        """FTS5 + jieba BM25 搜索"""
        try:
            import jieba
            segmented = " ".join(jieba.cut_for_search(query))
        except ImportError:
            segmented = query

        safe_query = self._sanitize_fts_query(segmented)
        conn = self._get_conn()
        result: dict[str, float] = {}

        try:
            if collection_id:
                rows = conn.execute(
                    "SELECT c.id, c.content, bm25(kb_chunks_fts) AS rank "
                    "FROM kb_chunks_fts f "
                    "JOIN kb_chunks c ON c.rowid = f.rowid "
                    "WHERE kb_chunks_fts MATCH ? AND c.collection_id = ? "
                    "ORDER BY rank LIMIT ?",
                    (safe_query, collection_id, limit),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT c.id, bm25(kb_chunks_fts) AS rank "
                    "FROM kb_chunks_fts f "
                    "JOIN kb_chunks c ON c.rowid = f.rowid "
                    "WHERE kb_chunks_fts MATCH ? "
                    "ORDER BY rank LIMIT ?",
                    (safe_query, limit),
                ).fetchall()
        except Exception:
            logger.debug("FTS query failed", exc_info=True)
            rows = []

        conn.close()
        # Convert BM25 rank to score: lower rank = better → score = 1/(1+rank)
        for r in rows:
            rank = r["rank"] if isinstance(r, sqlite3.Row) else r[1]
            chunk_id = r["id"] if isinstance(r, sqlite3.Row) else r[0]
            result[chunk_id] = 1.0 / (1.0 + rank)
        return result

    def _search_vector(
        self, query: str, collection_id: str | None, limit: int
    ) -> dict[str, float]:
        """ChromaDB 向量搜索"""
        if not self._vector_enabled or not self._chroma_collection:
            return {}

        try:
            embedding = self._embedding_model.encode(
                [query], show_progress_bar=False
            ).tolist()[0]

            where_filter = None
            if collection_id:
                where_filter = {"collection_id": collection_id}

            results = self._chroma_collection.query(
                query_embeddings=[embedding],
                n_results=limit,
                where=where_filter,
                include=["metadatas", "distances"],
            )

            scores: dict[str, float] = {}
            if results["ids"] and results["ids"][0]:
                ids_list = results["ids"][0]
                dists_list = results["distances"][0]
                mds_list = results.get("metadatas", [[]])[0] or []
                for idx, (cid, dist) in enumerate(zip(ids_list, dists_list)):
                    # metadatas[idx] is a dict: {"chunk_id": "...", "collection_id": "..."}
                    meta = mds_list[idx] if idx < len(mds_list) else {}
                    chunk_id = meta.get("chunk_id", "") if isinstance(meta, dict) else ""
                    key = chunk_id if chunk_id else cid
                    scores[key] = max(0, 1.0 - dist)
            return scores

        except Exception as e:
            logger.warning(f"[KnowledgeManager] 向量搜索失败: {e}")
            return {}

    @staticmethod
    def _rrf_fusion(
        bm25_scores: dict[str, float],
        vector_scores: dict[str, float],
        k: int = 60,
    ) -> list[SearchResult]:
        """Reciprocal Rank Fusion 融合 BM25 和向量排名"""
        def _ranked(scores: dict[str, float]) -> dict[str, int]:
            sorted_items = sorted(scores.items(), key=lambda x: x[1], reverse=True)
            return {chunk_id: rank + 1 for rank, (chunk_id, _) in enumerate(sorted_items)}

        bm25_ranks = _ranked(bm25_scores)
        vector_ranks = _ranked(vector_scores)

        all_chunk_ids = set(bm25_ranks.keys()) | set(vector_ranks.keys())
        results: list[SearchResult] = []
        from dataclasses import field as _field

        for chunk_id in all_chunk_ids:
            bm25_rank = bm25_ranks.get(chunk_id, len(bm25_ranks) + 1)
            vector_rank = vector_ranks.get(chunk_id, len(vector_ranks) + 1)
            rrf_score = 1.0 / (k + bm25_rank) + 1.0 / (k + vector_rank)

            source = "hybrid"
            if chunk_id in bm25_scores and chunk_id not in vector_scores:
                source = "bm25"
            elif chunk_id in vector_scores and chunk_id not in bm25_scores:
                source = "vector"

            # Create a minimal result (chunk details loaded by caller if needed)
            chunk = KnowledgeChunk(id=chunk_id)
            results.append(SearchResult(
                chunk=chunk,
                score=rrf_score,
                bm25_score=bm25_scores.get(chunk_id, 0),
                vector_score=vector_scores.get(chunk_id, 0),
                source_type=source,
            ))

        return results

    def search_with_chunks(
        self,
        query: str,
        collection_id: str | None = None,
        top_k: int = 10,
    ) -> list[SearchResult]:
        """搜索并填充完整的 chunk 内容"""
        results = self.search(query, collection_id, top_k)
        conn = self._get_conn()

        filled: list[SearchResult] = []
        for r in results:
            row = conn.execute(
                "SELECT * FROM kb_chunks WHERE id = ?", (r.chunk.id,)
            ).fetchone()
            if row:
                r.chunk = _chunk_from_row(row)
                filled.append(r)

        conn.close()
        return filled

    # ── Helpers ─────────────────────────────────────────────────────────────

    def _get_chunk_embedding_ids(self, doc_id: str) -> list[str]:
        conn = self._get_conn()
        rows = conn.execute(
            "SELECT embedding_id FROM kb_chunks WHERE doc_id = ? AND embedding_id IS NOT NULL",
            (doc_id,),
        ).fetchall()
        conn.close()
        return [r["embedding_id"] for r in rows]

    def _mark_doc_status(self, doc_id: str, status: DocStatus) -> None:
        conn = self._get_conn()
        conn.execute(
            "UPDATE kb_documents SET status = ?, updated_at = ? WHERE id = ?",
            (status.value, time.time(), doc_id),
        )
        conn.commit()
        conn.close()

    def _mark_doc_ready(self, doc_id: str, chunk_count: int) -> None:
        conn = self._get_conn()
        conn.execute(
            "UPDATE kb_documents SET status = ?, chunk_count = ?, error_message = NULL, "
            "updated_at = ? WHERE id = ?",
            (DocStatus.READY.value, chunk_count, time.time(), doc_id),
        )
        conn.commit()
        conn.close()

    def _mark_doc_error(self, doc_id: str, error: str) -> None:
        conn = self._get_conn()
        conn.execute(
            "UPDATE kb_documents SET status = ?, error_message = ?, updated_at = ? WHERE id = ?",
            (DocStatus.ERROR.value, error, time.time(), doc_id),
        )
        conn.commit()
        conn.close()


# ── SQLite row → dataclass converters ─────────────────────────────────────────

def _doc_from_row(r: sqlite3.Row | dict[str, Any]) -> KnowledgeDocument:
    if isinstance(r, sqlite3.Row):
        d = dict(r)
    else:
        d = r
    import json as _json
    metadata = {}
    if d.get("metadata_json"):
        try:
            metadata = _json.loads(d["metadata_json"])
        except Exception:
            pass
    return KnowledgeDocument(
        id=d["id"], collection_id=d["collection_id"],
        filename=d["filename"], file_path=d.get("file_path"),
        file_type=d["file_type"], file_size=d.get("file_size", 0),
        checksum=d.get("checksum", ""), status=d.get("status", "pending"),
        chunk_count=d.get("chunk_count", 0),
        error_message=d.get("error_message"),
        uploaded_by=d.get("uploaded_by", ""),
        created_at=d.get("created_at", 0.0), updated_at=d.get("updated_at", 0.0),
        metadata=metadata,
    )


def _chunk_from_row(r: sqlite3.Row | dict[str, Any]) -> KnowledgeChunk:
    if isinstance(r, sqlite3.Row):
        d = dict(r)
    else:
        d = r
    import json as _json
    metadata = {}
    if d.get("metadata_json"):
        try:
            metadata = _json.loads(d["metadata_json"])
        except Exception:
            pass
    return KnowledgeChunk(
        id=d["id"], doc_id=d["doc_id"], collection_id=d["collection_id"],
        content=d.get("content", ""), chunk_index=d.get("chunk_index", 0),
        token_count=d.get("token_count", 0),
        embedding_id=d.get("embedding_id"),
        metadata=metadata,
    )


# ── Utils ─────────────────────────────────────────────────────────────────────

def _sha256_file(file_path: Path) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()
